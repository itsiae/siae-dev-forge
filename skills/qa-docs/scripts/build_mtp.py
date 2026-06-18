#!/usr/bin/env python3
"""
build_mtp.py — Genera MTP .docx patchando il template SIAE (anchor-based).
Uso: python3 build_mtp.py <mtp_data.json> [output_dir]

Percorso: copia template → unzip → patch settings (ET) → repack
          → apri con python-docx → patch testo + SOSTITUISCE SmartArt con
            tabella 3 colonne scalabile → salva.

Il blocco Piattaforme/Piattaforme Modificate/Sistemi Impattati viene reso
come tabella Word (header blu, righe che crescono) anziché SmartArt,
così funziona correttamente con qualsiasi numero di voci.
"""

import sys
import json
import re
import shutil
import zipfile
import tempfile
from pathlib import Path
import xml.etree.ElementTree as ET

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERRORE: python-docx non installato. Esegui: pip3 install python-docx")
    sys.exit(1)

SKILL_DIR     = Path(__file__).parent.parent
TEMPLATE_PATH = SKILL_DIR / "assets" / "MTP_template.docx"

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def load_data(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def safe_slug(text, maxlen=40):
    slug = re.sub(r"[^\w\s-]", "", text or "").strip()
    slug = re.sub(r"\s+", "_", slug)
    return slug[:maxlen] or "documento"


# ---------------------------------------------------------------------------
# Settings patch (XML diretto)
# ---------------------------------------------------------------------------

def patch_settings(settings_xml_path):
    """Assicura <w:updateFields w:val='true'/> per ricalcolo TOC."""
    tree = ET.parse(str(settings_xml_path))
    root = tree.getroot()
    tag = f"{{{NS_W}}}updateFields"
    existing = root.find(tag)
    if existing is None:
        uf = ET.SubElement(root, tag)
        uf.set(f"{{{NS_W}}}val", "true")
    else:
        existing.set(f"{{{NS_W}}}val", "true")
    tree.write(str(settings_xml_path), xml_declaration=True, encoding="UTF-8")
    print("  settings.xml: updateFields=true")


# ---------------------------------------------------------------------------
# Tabella Piattaforme (sostituisce SmartArt)
# ---------------------------------------------------------------------------

def _make_table_row(texts, fill_hex=None, header=False):
    """
    Costruisce una <w:tr>.
    - header=True  → sfondo fill_hex, testo bianco grassetto, centrato
    - header=False → sfondo fill_hex (opzionale), simbolo ■ + testo, allineato a sinistra
    """
    CELL_W      = "3212"
    HEADER_TEXT = "FFFFFF"
    DATA_TEXT   = "1F3864"   # blu scuro per le voci

    tr = OxmlElement("w:tr")
    trPr = OxmlElement("w:trPr")
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   "480" if header else "380")
    trH.set(qn("w:hRule"), "atLeast")   # altezza minima — si espande con il testo
    trPr.append(trH)
    tr.append(trPr)

    for text in texts:
        tc   = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")

        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    CELL_W)
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

        if fill_hex:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill_hex)
            tcPr.append(shd)

        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)
        tc.append(tcPr)

        p   = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")

        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center" if header else "left")
        pPr.append(jc)

        # Indentazione sinistra per le voci (non per l'header)
        if not header:
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "80")
            pPr.append(ind)

        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "60")
        spacing.set(qn("w:after"),  "60")
        pPr.append(spacing)
        p.append(pPr)

        if text:
            r   = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")

            if header:
                b = OxmlElement("w:b")
                rPr.append(b)
                color_el = OxmlElement("w:color")
                color_el.set(qn("w:val"), HEADER_TEXT)
                rPr.append(color_el)
            else:
                color_el = OxmlElement("w:color")
                color_el.set(qn("w:val"), DATA_TEXT)
                rPr.append(color_el)

            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "20")   # 10pt
            rPr.append(sz)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), "20")
            rPr.append(szCs)

            r.append(rPr)
            t_el = OxmlElement("w:t")
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            # Checkbox ■ solo per le voci dati non vuote
            t_el.text = text if header else f"■  {text}"
            r.append(t_el)
            p.append(r)

        tc.append(p)
        tr.append(tc)
    return tr


def replace_smartart_with_table(doc, piattaforme, piattaforme_modificate, sistemi_impattati):
    """
    Trova il paragrafo che ospita lo SmartArt nel documento e lo sostituisce
    con una tabella 3 colonne scalabile:
      Piattaforme | Piattaforme Modificate | Sistemi Impattati

    La tabella ha header blu (#1F3864), testo bianco grassetto, righe alternate
    grigio chiaro (#F2F2F2) e si espande a qualsiasi numero di voci.
    """
    HEADER_FILL = "1F3864"
    ALT_FILL    = "F2F2F2"

    # Trova il paragrafo contenente lo SmartArt diagramma SIAE.
    # Strategia: cerca DOPO il paragrafo "Obiettivi del Progetto" un elemento
    # con graphicData URI che contiene "diagram" (SmartArt).
    # NON usare AlternateContent generico: anche il logo ha AlternateContent.
    DIAGRAM_URI = "http://schemas.openxmlformats.org/drawingml/2006/diagram"

    # Individua l'indice del paragrafo heading "Obiettivi del Progetto"
    heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "Titolo" in para.style.name and "obiettivi del progetto" in para.text.lower():
            heading_idx = i
            break

    search_from = (heading_idx + 1) if heading_idx is not None else 0

    smartart_para = None
    for para in doc.paragraphs[search_from:]:
        # Ferma la ricerca al prossimo heading (non andare oltre 1.1)
        if "Titolo" in para.style.name and para.text.strip():
            break
        el = para._element
        for child in el.iter():
            # SmartArt: graphicData con URI diagramma
            if child.tag.endswith("}graphicData"):
                if DIAGRAM_URI in child.get("uri", ""):
                    smartart_para = para
                    break
        if smartart_para:
            break

    if smartart_para is None:
        print("  AVVISO: paragrafo SmartArt diagramma non trovato in sezione 1.1 — tabella non inserita")
        return

    # Costruisce la tabella
    max_rows = max(len(piattaforme), len(piattaforme_modificate), len(sistemi_impattati), 1)

    tbl = OxmlElement("w:tbl")

    # Proprietà tabella
    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9638")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    # Bordi: linea orizzontale tra righe, verticali invisibili
    tblBorders = OxmlElement("w:tblBorders")
    for side, color, sz in [
        ("top",      "FFFFFF", "4"),
        ("left",     "FFFFFF", "4"),
        ("bottom",   "FFFFFF", "4"),
        ("right",    "FFFFFF", "4"),
        ("insideH",  "D9D9D9", "4"),   # separatore orizzontale leggero
        ("insideV",  "FFFFFF", "4"),   # nessun separatore verticale
    ]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # Griglia colonne (tre colonne uguali)
    tblGrid = OxmlElement("w:tblGrid")
    for _ in range(3):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), "3212")
        tblGrid.append(col)
    tbl.append(tblGrid)

    # Riga header
    tbl.append(_make_table_row(
        ["Piattaforme", "Piattaforme Modificate", "Sistemi Impattati"],
        fill_hex=HEADER_FILL,
        header=True,
    ))

    # Righe dati
    for i in range(max_rows):
        p_val  = piattaforme[i]             if i < len(piattaforme)             else ""
        pm_val = piattaforme_modificate[i]  if i < len(piattaforme_modificate)  else ""
        si_val = sistemi_impattati[i]       if i < len(sistemi_impattati)        else ""
        fill   = ALT_FILL if i % 2 == 0 else None
        tbl.append(_make_table_row([p_val, pm_val, si_val], fill_hex=fill))

    # Sostituisce il paragrafo SmartArt con la tabella nel body XML
    parent  = smartart_para._element.getparent()
    sa_pos  = list(parent).index(smartart_para._element)
    parent.remove(smartart_para._element)
    parent.insert(sa_pos, tbl)

    print(
        f"  SmartArt → Tabella 3 colonne: "
        f"{len(piattaforme)}P + {len(piattaforme_modificate)}PM + "
        f"{len(sistemi_impattati)}SI = {max_rows} righe dati"
    )


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def find_heading_paragraph(doc, heading_text, style_contains="Titolo"):
    matches = []
    for i, para in enumerate(doc.paragraphs):
        if style_contains in para.style.name and heading_text.lower() in para.text.lower():
            matches.append(i)
    if not matches:
        raise ValueError(
            f"ANCHOR NON TROVATO: heading '{heading_text}' (stile ~'{style_contains}'). "
            f"Il template potrebbe essere diverso da quello atteso."
        )
    if len(matches) > 1:
        raise ValueError(
            f"ANCHOR AMBIGUO: '{heading_text}' trovato {len(matches)} volte a indici {matches}."
        )
    return matches[0]


def paragraphs_until_next_heading(doc, start_idx, stop_styles=None):
    if stop_styles is None:
        stop_styles = ("Titolo1", "Titolo2", "Titolo3", "Titolo4")
    result = []
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if any(s in doc.paragraphs[i].style.name for s in stop_styles):
            break
        result.append(i)
    return result


def set_para_text(para, text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


# ---------------------------------------------------------------------------
# Document patching
# ---------------------------------------------------------------------------

def patch_document(doc, data):
    meta = data.get("meta", {})

    # --- COPERTINA: codice DMND ---
    for para in doc.paragraphs[:20]:
        if re.search(r"DMND\d+", para.text.strip(), re.IGNORECASE):
            set_para_text(para, meta.get("codice", ""))
            break

    # --- COPERTINA: titolo ---
    found_dmnd = False
    for para in doc.paragraphs[:25]:
        text = para.text.strip()
        if re.search(r"DMND\d+", text, re.IGNORECASE):
            found_dmnd = True
            continue
        if found_dmnd and text and "Master Test Plan" not in text:
            set_para_text(para, meta.get("titolo", ""))
            break

    # --- TABELLA ANAGRAFICA: Compilato + Versione ---
    for table in doc.tables[:4]:
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                ct = cell.text.strip()
                if ct == "Compilato" and ci + 1 < len(row.cells):
                    target = row.cells[ci + 1].paragraphs[0]
                    if target.runs:
                        target.runs[0].text = meta.get("autore", "")
                    else:
                        target.add_run(meta.get("autore", ""))
                if ct == "Versione" and ci + 1 < len(row.cells):
                    target = row.cells[ci + 1].paragraphs[0]
                    if target.runs:
                        target.runs[0].text = meta.get("versione", "1.0")
                    else:
                        target.add_run(meta.get("versione", "1.0"))

    # --- REGISTRO MODIFICHE: prima riga dati ---
    for table in doc.tables[:6]:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        if any(h in headers for h in ("Vrs.", "Data", "Autore")):
            if len(table.rows) > 1:
                row = table.rows[1]
                vals = [
                    meta.get("versione", "1.0"),
                    meta.get("data", ""),
                    "Tutti",
                    meta.get("autore", ""),
                    "First release",
                ]
                for ci, val in enumerate(vals):
                    if ci < len(row.cells):
                        p = row.cells[ci].paragraphs[0]
                        if p.runs:
                            p.runs[0].text = val
                        else:
                            p.add_run(val)
            break

    # --- 1.1 Obiettivi + SOSTITUZIONE SmartArt con tabella ---
    try:
        idx = find_heading_paragraph(doc, "Obiettivi del Progetto")
        to_patch = paragraphs_until_next_heading(doc, idx)
        obiettivo = data.get("obiettivo_progetto", "")
        if to_patch:
            set_para_text(doc.paragraphs[to_patch[0]], obiettivo)
            for extra in to_patch[1:]:
                p = doc.paragraphs[extra]
                if "Di seguito" in p.text:
                    continue
                if p.text.strip() and not any(
                    elem.tag.endswith("}drawing") or elem.tag.endswith("}object")
                    or "AlternateContent" in elem.tag
                    for elem in p._element.iter()
                ):
                    set_para_text(p, "")
    except ValueError as e:
        print(f"  AVVISO: {e}")

    # Sostituisci SmartArt con tabella scalabile
    replace_smartart_with_table(
        doc,
        data.get("piattaforme", []),
        data.get("piattaforme_modificate", []),
        data.get("sistemi_impattati", []),
    )

    # --- 1.3 Tabella livelli ---
    for table in doc.tables:
        if table.rows and table.rows[0].cells and "Livello" in table.rows[0].cells[0].text:
            livelli = data.get("livelli", [])
            for li, lv in enumerate(livelli):
                ri = li + 1
                if ri < len(table.rows):
                    row = table.rows[ri]
                    vals = [
                        lv.get("livello", ""),
                        lv.get("descrizione", ""),
                        lv.get("dettaglio", ""),
                        lv.get("owner", ""),
                        lv.get("governance", ""),
                    ]
                    for ci, val in enumerate(vals):
                        if ci < len(row.cells):
                            p = row.cells[ci].paragraphs[0]
                            if p.runs:
                                p.runs[0].text = val
                            else:
                                p.add_run(val)
            break

    # --- Sezioni testo semplice ---
    simple_sections = [
        ("Performance test", data.get("performance_test", "Non previsti.")),
        ("NRT",              data.get("nrt",              "Non previsti.")),
        ("Test Automatici",  data.get("test_automatici",  "Non previsti.")),
    ]
    for heading, text in simple_sections:
        try:
            idx = find_heading_paragraph(doc, heading)
            to_patch = paragraphs_until_next_heading(doc, idx)
            if to_patch:
                set_para_text(doc.paragraphs[to_patch[0]], text)
                for extra in to_patch[1:]:
                    p = doc.paragraphs[extra]
                    if p.text.strip():
                        set_para_text(p, "")
        except ValueError as e:
            print(f"  AVVISO: {e}")

    # --- GANTT ---
    try:
        idx = find_heading_paragraph(doc, "GANTT")
        to_patch = paragraphs_until_next_heading(doc, idx)
        gantt = data.get("gantt", [])
        gantt_text = (
            "\n".join(
                f"Sprint {g.get('sprint','?')}: {g.get('date','?')} — {g.get('pct','?')}%"
                for g in gantt
            )
            if gantt else "DA CONFERMARE"
        )
        if to_patch:
            set_para_text(doc.paragraphs[to_patch[0]], gantt_text)
    except ValueError as e:
        print(f"  AVVISO: {e}")

    # --- Out of Scope ---
    try:
        idx = find_heading_paragraph(doc, "Out of Scope")
        to_patch = paragraphs_until_next_heading(doc, idx)
        for i, item in enumerate(data.get("out_of_scope", [])):
            if i < len(to_patch):
                set_para_text(doc.paragraphs[to_patch[i]], f"• {item}")
    except ValueError as e:
        print(f"  AVVISO: {e}")

    # --- Rischi ---
    try:
        idx = find_heading_paragraph(doc, "Rischi, Problemi")
        to_patch = paragraphs_until_next_heading(doc, idx)
        for i, item in enumerate(data.get("rischi", [])):
            if i < len(to_patch):
                set_para_text(doc.paragraphs[to_patch[i]], f"• {item}")
    except ValueError as e:
        print(f"  AVVISO: {e}")

    print("  document.xml patchato")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build_mtp(json_path, output_dir):
    data      = load_data(json_path)
    meta      = data.get("meta", {})
    codice    = meta.get("codice", "MTP")
    titolo    = safe_slug(meta.get("titolo", "documento"))
    out_name  = f"MTP_{codice}_-_{titolo}.docx"

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / out_name

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template non trovato: {TEMPLATE_PATH}")

    print(f"Template: {TEMPLATE_PATH}")
    print(f"Output:   {output_path}")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        tmp_docx = tmpdir / "working.docx"
        shutil.copy2(str(TEMPLATE_PATH), str(tmp_docx))

        # Unzip per patch settings
        xml_dir = tmpdir / "xml"
        with zipfile.ZipFile(str(tmp_docx), "r") as z:
            z.extractall(str(xml_dir))

        settings_path = xml_dir / "word" / "settings.xml"
        if settings_path.exists():
            patch_settings(settings_path)

        # Repack
        repacked = tmpdir / "repacked.docx"
        with zipfile.ZipFile(str(repacked), "w", zipfile.ZIP_DEFLATED) as z:
            for item in xml_dir.rglob("*"):
                if item.is_file():
                    z.write(str(item), str(item.relative_to(xml_dir)))

        # Patch con python-docx (testo + tabella piattaforme)
        doc = Document(str(repacked))
        patch_document(doc, data)
        doc.save(str(output_path))

    print(f"\nGenerato: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(f"Uso: {sys.argv[0]} <mtp_data.json> [output_dir]")
        sys.exit(1)

    json_path = sys.argv[1]
    out_dir   = sys.argv[2] if len(sys.argv) > 2 else "."

    try:
        result = build_mtp(json_path, out_dir)
        print(f"OK: {result}")
    except Exception as e:
        print(f"ERRORE: {e}")
        import traceback; traceback.print_exc()
        sys.exit(1)
